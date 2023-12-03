from azure.core.credentials import AzureKeyCredential
from azure.ai.language.questionanswering import QuestionAnsweringClient
from azure.ai.language.questionanswering import models as qna
import streamlit as st

key = st.secrets['LANGUAGE_KEY']
endpoint = st.secrets['CHATBOT_ENDPOINT']

credential = AzureKeyCredential(key)

st.set_page_config( 
     page_title="CogniQuotient", 
     page_icon="🏫",
     initial_sidebar_state="expanded", 
 ) 


if 'username' not in st.session_state:
    st.session_state["username"]="User"
    
import docx

def bot_response(knowledge, question):
    client = QuestionAnsweringClient(endpoint, credential)
    text_documents = knowledge.split('\n')
    print(knowledge)
    with client:
        input = qna.AnswersFromTextOptions(
            question=question,
            text_documents= text_documents
        )
        output = client.get_answers_from_text(input)
        answer = False
        best_answer = {'answer':"Please enter more information."}
    try:
        best_answer = [a for a in output.answers if a.confidence > 0.83][0]
        answer = True
    except:
        best_answer={'answer':"Please enter more information."}
    print(u"Q: {}".format(input.question))
    if answer:
        return(u"{}".format(best_answer.answer))
    else:
        return(u"{}".format(best_answer['answer']))

def save_to_doc(conversation):
    doc = docx.Document()
    doc.add_heading("Chatbot Conversation", level=1)

    for user, bot in conversation:
        p_user = doc.add_paragraph()
        p_user.add_run("User: ").bold = True
        p_user.add_run(user)

        p_bot = doc.add_paragraph()
        p_bot.add_run("Bot: ").bold = True
        p_bot.add_run(bot)

    doc.save("chatbot_conversation.docx")

def runner():
    
    if not knowledge:
        st.stop()

    if "conversation" not in st.session_state:
        st.session_state.conversation = []

    user_input = st.chat_input(placeholder="Your message")

    if user_input:
        if user_input:
            bot_reply = bot_response(st.session_state.knowledge,user_input)
            st.session_state.conversation.append((user_input, bot_reply))
        else:
            st.warning("Please enter a question.")
        
        save_to_doc(st.session_state.conversation)

    st.text("Conversation History:")
    for user, bot in st.session_state.conversation:
        text = st.chat_message("User")
        message = st.chat_message("Assistant")
        text.write(user)
        message.write(bot)
    
    if st.button("End Conversation"):
        st.session_state.conversation=[]
        st.session_state.knowledge=""
        st.session_state.first_run = True
        st.snow()
        st.success("Click on any button to refresh", icon='✅')
        return
    
    st.download_button(
                label="Download Conversation",
                data=open("chatbot_conversation.docx", "rb").read(),
                file_name="Conversation.docx",
                mime="application/octet-stream",
                help="Click to download the conversation."
            ) 


def gradient_text(text, color1, color2):
    gradient_css = f"""
        background: -webkit-linear-gradient(left, {color1}, {color2});
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-weight: bold;
        font-size: 42px;
    """
    return f'<span style="{gradient_css}">{text}</span>'

def gradient(text, color1, color2):
    gradient_css = f"""
        background: -webkit-linear-gradient(left, {color2}, {color1});
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-style: italic;
        font-size: 22px;
    """
    return f'<span style="{gradient_css}">{text}</span>'


color1 = "#0d3270"
color2 = "#0fab7b"
text = "Chat & Learn"

styled_text = gradient_text(text, color1, color2)
st.write(f"<div style='text-align: center;'>{styled_text}</div>", unsafe_allow_html=True)
styled_text = gradient("Explore, Learn while chatting about yout text", color1, color2)
st.write(f"<div style='text-align: center;'>{styled_text}</div>", unsafe_allow_html=True)
st.subheader("Weclome "+st.session_state["username"]+" 👋")

first_run = st.session_state.get("first_run", True)

if first_run:
    knowledge = st.text_input("Please the text you want to query:")
    if st.button('Upload'):
        st.session_state.knowledge = knowledge
        st.session_state.first_run = False
else:
    knowledge = st.session_state.get("knowledge", "")
    runner()

footer="""<style>

a:hover,  a:active {
color: red;
background-color: transparent;
text-decoration: underline;
}

.footer {
position: fixed;
left: 0;
bottom: 0;
width: 100%;
background-color: white;
color: black;
text-align: center;
}
</style>
<div class="footer">
<p>Developed with ❤️ by <a style='display: inline; text-align: center;' href="https://www.linkedin.com/in/harshavardhan-bajoria/" target="_blank">Harshavardhan Bajoria</a></p>
</div>
"""
st.markdown(footer,unsafe_allow_html=True)

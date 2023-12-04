import streamlit as st
from azure.ai.textanalytics import TextAnalyticsClient
from azure.core.credentials import AzureKeyCredential
from docx import Document

key = os.environ['LANGUAGE_KEY']
endpoint = os.environ['LANGUAGE_ENDPOINT']

credential = AzureKeyCredential(key)
text_analytics_client = TextAnalyticsClient(endpoint, credential)

st.set_page_config( 
     page_title="CogniQuotient", 
     page_icon="🏫",
     initial_sidebar_state="expanded", 
 ) 

if 'username' not in st.session_state:
    st.session_state["username"]="User"
    
# Streamlit app
def main():

    def gradient_text(text, color1, color2):
        gradient_css = f"""
        background: -webkit-linear-gradient(left, {color1}, {color2});
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-weight: bold;
        font-size: 42px;
        """
        return f'<span style="{gradient_css}">{text}</span>'
    
    def gradient(text, color1, color2):
        gradient_css = f"""
        background: -webkit-linear-gradient(left, {color2}, {color1});
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-style: italic;
        font-size: 22px;
        """
        return f'<span style="{gradient_css}">{text}</span>'


    color1 = "#0d3270"
    color2 = "#0fab7b"
    text = "Keyword-Powered Notes"

    styled_text = gradient_text(text, color1, color2)
    st.write(f"<div style='text-align: center;'>{styled_text}</div>", unsafe_allow_html=True)
    styled_text = gradient("Extract, Highlight Important Keywords", color1, color2)
    st.write(f"<div style='text-align: center;'>{styled_text}</div>", unsafe_allow_html=True)
    st.subheader("Weclome "+st.session_state["username"]+" 👋")
   
    text_input = st.text_area("Enter your text here:")
    if st.button("Extract Keywords"):
        if text_input:
            keywords = extract_keywords(text_input)
            keywords = sort_keywords_by_occurrence(keywords, text_input)
            st.success("Here's Your Note! Download it below.",icon='✅')
            st.subheader("Extracted Keywords:")
            st.write(", ".join(keywords))
            note_text = f"{text_input}"
            doc = create_doc_with_bold_keywords(note_text, keywords)
            doc.save("note.docx")
            
            st.download_button(
                label="Download Note",
                data=open("note.docx", "rb").read(),
                file_name="note.docx",
                mime="application/octet-stream",
                help="Click to download the note with bold keywords."
            )      

def extract_keywords(text):
    response = text_analytics_client.extract_key_phrases(documents=[text])
    return response[0].key_phrases

def sort_keywords_by_occurrence(keywords, text):
    keyword_indices = {keyword: text.lower().find(keyword.lower()) for keyword in keywords}
    return sorted(keywords, key=lambda keyword: keyword_indices[keyword])

def create_doc_with_bold_keywords(note_text, keywords):
    doc = Document()
    doc.add_heading('Here\'s The Note ', level=1)
    p = doc.add_paragraph()
    start_idx = 0
    for keyword in keywords:
        keyword_start = note_text.find(keyword, start_idx)
        if keyword_start != -1:
            keyword_end = keyword_start + len(keyword)
            p.add_run(note_text[start_idx:keyword_start])
            p.add_run(note_text[keyword_start:keyword_end]).bold = True
            start_idx = keyword_end
    p.add_run(note_text[start_idx:])
    return doc

if __name__ == "__main__":
    main()
footer="""<style>

a:hover,  a:active {
color: red;
background-color: transparent;
text-decoration: underline;
}

.footer {
position: fixed;
left: 0;
bottom: 0;
width: 100%;
background-color: white;
color: black;
text-align: center;
}
</style>
<div class="footer">
<p>Developed with ❤️ by <a style='display: inline; text-align: center;' href="https://www.linkedin.com/in/harshavardhan-bajoria/" target="_blank">Harshavardhan Bajoria</a></p>
</div>
"""
st.markdown(footer,unsafe_allow_html=True)
